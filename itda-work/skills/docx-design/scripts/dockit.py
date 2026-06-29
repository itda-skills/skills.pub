# -*- coding: utf-8 -*-
"""dockit — docx-design 공개 헬퍼 API (python-docx 래퍼).

pptx-design 의 deckkit 에 대응하는 docx 어댑터다. 다만 pptx 가 16:9 캔버스에
**절대좌표**로 도형을 쌓는 것과 달리, docx 는 **흐름(flow) 문서**라 단락·헤딩·표·
헤더/푸터가 페이지를 따라 흐른다. 따라서 본 모듈은 좌표 헬퍼가 아니라 "디자인된
단락/표/콜아웃/규칙선/푸터" 빌더를 제공한다.

★한글 East Asian 바인딩(SPEC-OFFICE-DOC-GEN-DEEPEN-001 — "docx 맥락 재설계"):
  pptx 의 cjk_guard 는 LibreOffice 의 한글 세리프/붓글씨 폴백을 막으려는 것이지만,
  Word 는 한글을 네이티브로 정상 렌더한다. docx 의 핵심은 run 의 `w:rFonts` 를
  ascii/hAnsi(라틴) ↔ eastAsia(한글)로 **분리 바인딩**하는 정확성이다 — 라틴 디스플레이
  폰트가 한글 글리프를 대신 먹어 깨지는 일을 차단한다. 음수 자간 클램프는 불필요
  (Word 가 자간을 정상 처리)하므로 라틴 트래킹은 그대로 허용한다.

백엔드: python-docx(크로스플랫폼, Office 불필요)가 1급. 토큰은 design-core 의
`docx_styles()` 로부터 받는다(매체 중립 SSOT).
"""
from __future__ import annotations

import os
import platform
import re

from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────────────── 색/폰트 유틸
_HEX_RE = re.compile(r"^#?[0-9a-fA-F]{3}([0-9a-fA-F]{3})?$")


def _hex6(value) -> str:
    """'#1E2761' / '1e2761' / '#abc' → 'RRGGBB'(대문자 6자리)."""
    h = str(value or "").strip().lstrip("#")
    if len(h) == 3:
        h = "".join(c * 2 for c in h)
    return h.upper()


def rgb(value) -> RGBColor:
    return RGBColor.from_string(_hex6(value))


def luminance(value) -> float:
    h = _hex6(value)
    r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
    return 0.299 * r + 0.587 * g + 0.114 * b


def on_color(value, dark="#1A1A1A", light="#FFFFFF") -> str:
    """배경색 위에 올릴 글자색을 명도로 자동 선택(어두우면 흰색, 밝으면 잉크)."""
    return light if luminance(value) < 140 else dark


def has_hangul(text) -> bool:
    return any("가" <= ch <= "힣" or "㄰" <= ch <= "㆏" for ch in (text or ""))


# eastAsia 후보 → 폰트 파일 힌트(존재 탐지용)
_KR_FILE_HINTS = {
    "Malgun Gothic": ["malgun.ttf", "malgunbd.ttf"],
    "맑은 고딕": ["malgun.ttf"],
    "Noto Sans KR": ["NotoSansKR-Regular.otf", "NotoSansKR-Regular.ttf",
                     "NotoSansCJKkr-Regular.otf", "NotoSansKR[wght].ttf"],
    "Pretendard": ["Pretendard-Regular.ttf", "Pretendard-Regular.otf",
                   "PretendardVariable.ttf"],
    "AppleSDGothicNeo": ["AppleSDGothicNeo.ttc"],
    "나눔고딕": ["NanumGothic.ttf"],
    "NanumGothic": ["NanumGothic.ttf"],
}

_DEFAULT_KR_CANDIDATES = ["Malgun Gothic", "맑은 고딕", "Noto Sans KR", "Pretendard"]


def _font_dirs():
    sysname = platform.system()
    dirs = []
    if sysname == "Windows":
        windir = os.environ.get("WINDIR", r"C:\Windows")
        dirs += [os.path.join(windir, "Fonts"),
                 os.path.join(os.environ.get("LOCALAPPDATA", ""),
                              "Microsoft", "Windows", "Fonts")]
    elif sysname == "Darwin":
        dirs += ["/System/Library/Fonts", "/Library/Fonts",
                 os.path.expanduser("~/Library/Fonts")]
    else:
        dirs += ["/usr/share/fonts", "/usr/local/share/fonts",
                 os.path.expanduser("~/.fonts"),
                 os.path.expanduser("~/.local/share/fonts")]
    return [d for d in dirs if d and os.path.isdir(d)]


def _font_installed(name) -> bool:
    hints = _KR_FILE_HINTS.get(name)
    if not hints:
        return False
    for d in _font_dirs():
        for h in hints:
            if os.path.exists(os.path.join(d, h)):
                return True
    return False


def kr_font_name(candidates=None) -> str:
    """이 환경에서 Word/LibreOffice 가 한글을 또렷하게 렌더하는 eastAsia 폰트명.

    우선순위: 환경변수 DOCX_KR_FONT → candidates(존재하는 첫 폰트) → 플랫폼 기본.
    Windows 는 Malgun Gothic 이 사실상 항상 존재(Vista+)하므로 폴백도 그것.
    """
    env = os.environ.get("DOCX_KR_FONT")
    if env:
        return env
    cands = list(candidates or _DEFAULT_KR_CANDIDATES)
    for c in cands:
        if _font_installed(c):
            return c
    if platform.system() == "Windows":
        return "Malgun Gothic"
    return cands[0] if cands else "Noto Sans KR"


# ──────────────────────────────────────────────────────── run 폰트(★eastAsia 분리 바인딩)
def set_run_font(run, latin=None, kr=None, size=None, bold=None, italic=None,
                 color=None, spacing=None, all_caps=None, force_kr=False):
    """run 폰트 설정 — ★라틴(ascii/hAnsi)과 한글(eastAsia)을 분리 바인딩.

    - latin: 라틴/숫자 글리프 폰트(`w:ascii`·`w:hAnsi`).
    - kr: 한글 글리프 폰트(`w:eastAsia`). None 이고 run 에 한글이 있으면(또는 force_kr)
      kr_font_name() 안전 고딕을 자동 바인딩한다.
    - spacing: 라틴 트래킹(pt, 음수 허용 — Word 는 정상 처리). 한글에도 그대로 둔다.
    """
    f = run.font
    if size is not None:
        f.size = Pt(float(size))
    if bold is not None:
        f.bold = bold
    if italic is not None:
        f.italic = italic
    if color is not None:
        f.color.rgb = rgb(color)
    if all_caps is not None:
        f.all_caps = all_caps

    want_kr = kr if kr is not None else (kr_font_name() if (force_kr or has_hangul(run.text)) else None)
    if latin is not None or want_kr is not None:
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        if latin is not None:
            rfonts.set(qn("w:ascii"), latin)
            rfonts.set(qn("w:hAnsi"), latin)
        if want_kr is not None:
            rfonts.set(qn("w:eastAsia"), want_kr)

    if spacing is not None:
        rpr = run._element.get_or_add_rPr()
        sp = rpr.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            # 스키마 순서: color 다음, w/kern/sz 앞
            rpr.insert_element_before(
                sp, "w:w", "w:kern", "w:position", "w:sz", "w:szCs", "w:highlight",
                "w:u", "w:effect", "w:bdr", "w:shd", "w:fitText", "w:vertAlign",
                "w:rtl", "w:cs", "w:em", "w:lang", "w:eastAsianLayout",
                "w:specVanish", "w:oMath",
            )
        sp.set(qn("w:val"), str(int(round(float(spacing) * 20))))  # pt → 1/20 pt


# ──────────────────────────────────────────────────────── 정렬 맵
_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT, "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
_VALIGN = {
    "top": WD_ALIGN_VERTICAL.TOP, "center": WD_ALIGN_VERTICAL.CENTER,
    "bottom": WD_ALIGN_VERTICAL.BOTTOM,
}


def _norm_runs(runs):
    """문자열 또는 [(text, fontkwargs), ...] 를 표준 run 리스트로."""
    if runs is None:
        return []
    if isinstance(runs, str):
        return [(runs, {})]
    return list(runs)


# ──────────────────────────────────────────────────────── 문서 생성/디자인 적용
def new_doc(margin_in=1.0, page="A4") -> Document:
    doc = Document()
    sec = doc.sections[0]
    if str(page).lower() == "letter":
        sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    else:  # A4 기본
        sec.page_width, sec.page_height = Mm(210), Mm(297)
    m = Inches(float(margin_in))
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = m
    return doc


def _style_set_fonts(style, latin=None, kr=None):
    """단락 스타일 rPr 의 rFonts 를 라틴/한글 분리 지정."""
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    if latin:
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
    if kr:
        rfonts.set(qn("w:eastAsia"), kr)


def apply_design(doc, st: dict) -> dict:
    """design-core docx_styles() dict 로 Normal·Heading 스타일·여백을 설정하고 theme 반환.

    theme 는 doc 에 stash 되어 이후 빌더(heading/body/table/footer)가 기본 색·폰트로 읽는다.
    """
    latin = st.get("latin_font") or "Calibri"
    kr = kr_font_name(st.get("korean_fallbacks"))
    ink = st.get("ink") or "#222222"
    primary = st.get("primary") or ink
    # 헤딩은 라이트 본문 위 텍스트 — 브랜드색이 저대비면 design-core 가 보정한
    # primary_text 를 쓴다(대비 충족 프리셋은 primary 와 동일).
    head = st.get("primary_text") or primary
    ts = st.get("type_scale") or {}

    normal = doc.styles["Normal"]
    normal.font.name = latin
    normal.font.size = Pt(float(ts.get("body", 10.5)))
    normal.font.color.rgb = rgb(ink)
    _style_set_fonts(normal, latin, kr)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.3

    for lvl, size_def, col in ((1, 22, head), (2, 15, head), (3, 12.5, ink)):
        s = doc.styles[f"Heading {lvl}"]
        s.font.name = latin
        s.font.bold = True
        s.font.size = Pt(float(ts.get(f"h{lvl}", size_def)))
        s.font.color.rgb = rgb(col)
        _style_set_fonts(s, latin, kr)
        s.paragraph_format.space_before = Pt(16 if lvl == 1 else 11)
        s.paragraph_format.space_after = Pt(4)
        s.paragraph_format.keep_with_next = True

    theme = dict(st)
    theme["_latin"], theme["_kr"] = latin, kr
    doc._dockit_theme = theme  # type: ignore[attr-defined]
    return theme


def theme_of(doc) -> dict:
    return getattr(doc, "_dockit_theme", {}) or {}


def save_doc(doc, path):
    os.makedirs(os.path.dirname(os.path.abspath(path)), exist_ok=True)
    doc.save(path)
    return path


# ──────────────────────────────────────────────────────── 단락 빌더
def add_paragraph(container, runs, style=None, align=None, space_before=None,
                  space_after=None, line_spacing=None):
    """container(doc 또는 cell)에 디자인 단락 추가. runs: str 또는 [(text, fontkwargs)]."""
    p = container.add_paragraph(style=style) if style else container.add_paragraph()
    if align is not None:
        p.alignment = _ALIGN[align]
    pf = p.paragraph_format
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    for text, kw in _norm_runs(runs):
        r = p.add_run(text)
        set_run_font(r, **kw)
    return p


def _with_theme_fonts(doc, runs):
    th = theme_of(doc)
    out = []
    for text, kw in _norm_runs(runs):
        kw = dict(kw)
        kw.setdefault("latin", th.get("_latin"))
        kw.setdefault("kr", th.get("_kr"))
        out.append((text, kw))
    return out


def heading(doc, runs, level=1, color=None, align=None, space_before=None, space_after=None):
    th = theme_of(doc)
    p = doc.add_paragraph(style=f"Heading {level}")
    if align is not None:
        p.alignment = _ALIGN[align]
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    for text, kw in _norm_runs(runs):
        kw = dict(kw)
        kw.setdefault("latin", th.get("_latin"))
        kw.setdefault("kr", th.get("_kr"))
        if color:
            kw.setdefault("color", color)
        r = p.add_run(text)
        set_run_font(r, **kw)
    return p


def kicker(doc, text, color=None, size=9, space_after=2):
    """소형 대문자 트래킹 라벨(섹션 킥커). 라틴 트래킹은 Word 가 정상 처리."""
    th = theme_of(doc)
    col = color or th.get("muted") or "#777777"
    return add_paragraph(
        doc, [(text, dict(latin=th.get("_latin"), size=size, color=col,
                          bold=True, all_caps=True, spacing=1.2))],
        space_after=space_after,
    )


def body(doc, runs, align=None, space_before=None, space_after=None, line_spacing=None):
    return add_paragraph(doc, _with_theme_fonts(doc, runs), align=align,
                         space_before=space_before, space_after=space_after,
                         line_spacing=line_spacing)


def bullet_list(doc, items, color=None):
    th = theme_of(doc)
    out = []
    for it in items:
        runs = []
        for text, kw in _norm_runs(it):
            kw = dict(kw)
            kw.setdefault("latin", th.get("_latin"))
            kw.setdefault("kr", th.get("_kr"))
            if color:
                kw.setdefault("color", color)
            runs.append((text, kw))
        p = doc.add_paragraph(style="List Bullet")
        for text, kw in runs:
            r = p.add_run(text)
            set_run_font(r, **kw)
        out.append(p)
    return out


def _para_border(paragraph, edge, color, weight_pt=1.0, space=1):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    el = pbdr.find(qn("w:" + edge))
    if el is None:
        el = OxmlElement("w:" + edge)
        pbdr.append(el)
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(int(round(weight_pt * 8))))  # 1/8 pt
    el.set(qn("w:space"), str(space))
    el.set(qn("w:color"), _hex6(color))


def rule(doc, color=None, weight_pt=1.0, space_before=4, space_after=10):
    """수평 규칙선(빈 단락의 하단 테두리)."""
    th = theme_of(doc)
    col = color or th.get("border") or "#CCCCCC"
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.0
    _para_border(p, "bottom", col, weight_pt)
    return p


# ──────────────────────────────────────────────────────── 표/셀 스타일
def shade_cell(cell, fill):
    if not fill:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _hex6(fill))


def set_cell_margins(cell, top=None, left=None, bottom=None, right=None):
    """셀 안쪽 여백(dxa = 1/20 pt). 가독 패딩용."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for name, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        if val is None:
            continue
        el = tcMar.find(qn("w:" + name))
        if el is None:
            el = OxmlElement("w:" + name)
            tcMar.append(el)
        el.set(qn("w:w"), str(int(val)))
        el.set(qn("w:type"), "dxa")


def set_cell_border(cell, **edges):
    """edges: top/left/bottom/right = dict(val,sz,color,space). 액센트 바·헤어라인용."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, spec in edges.items():
        el = tcBorders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            tcBorders.append(el)
        el.set(qn("w:val"), spec.get("val", "single"))
        el.set(qn("w:sz"), str(int(spec.get("sz", 8))))
        el.set(qn("w:space"), str(int(spec.get("space", 0))))
        el.set(qn("w:color"), _hex6(spec.get("color", "000000")))


def set_table_borders(table, color, weight_pt=0.75, inside=True, outside=True):
    col = _hex6(color)
    sz = str(int(round(weight_pt * 8)))
    tblPr = table._tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    edges = []
    if outside:
        edges += ["top", "left", "bottom", "right"]
    if inside:
        edges += ["insideH", "insideV"]
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        if edge in edges:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), col)
        else:
            el.set(qn("w:val"), "none")


def set_cell_runs(cell, runs, latin=None, kr=None, bold=None, color=None, size=None,
                  align=None, vertical="center", first=True):
    """셀의 (첫) 단락에 run 추가. first=False 면 새 단락."""
    p = cell.paragraphs[0] if (first and cell.paragraphs) else cell.add_paragraph()
    if align is not None:
        p.alignment = _ALIGN[align]
    if vertical is not None:
        cell.vertical_alignment = _VALIGN.get(vertical, WD_ALIGN_VERTICAL.CENTER)
    for text, kw in _norm_runs(runs):
        kw = dict(kw)
        if latin is not None:
            kw.setdefault("latin", latin)
        if kr is not None:
            kw.setdefault("kr", kr)
        if bold is not None:
            kw.setdefault("bold", bold)
        if color is not None:
            kw.setdefault("color", color)
        if size is not None:
            kw.setdefault("size", size)
        r = p.add_run(text)
        set_run_font(r, **kw)
    return p


def add_table(doc, headers, rows, header_fill=None, header_color=None,
              zebra=True, zebra_fill=None, border_color=None, col_align=None,
              header_align=None, font_size=10, header_size=10.5, width_in=None,
              col_widths=None):
    """디자인된 데이터 표: 헤더 fill + zebra + 헤어라인 테두리 + 한글 eastAsia 바인딩.

    headers/rows 셀은 str 또는 [(text, fontkwargs)]. col_align: 열별 정렬 리스트.
    """
    th = theme_of(doc)
    latin, kr = th.get("_latin"), th.get("_kr")
    header_fill = header_fill or th.get("accent") or th.get("primary") or "#333333"
    header_color = header_color or on_color(header_fill)
    zebra_fill = zebra_fill or th.get("surface") or "#F2F2F2"
    border_color = border_color or th.get("border") or "#CCCCCC"
    ncol = len(headers)
    aligns = col_align or (["left"] + ["right"] * (ncol - 1) if ncol > 1 else ["left"])
    halign = header_align or aligns

    table = doc.add_table(rows=1 + len(rows), cols=ncol)
    table.style = None
    table.autofit = False if (width_in or col_widths) else True
    set_table_borders(table, border_color, weight_pt=0.75, inside=True, outside=True)

    # 헤더
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        shade_cell(c, header_fill)
        set_cell_margins(c, top=70, bottom=70, left=110, right=110)
        set_cell_runs(c, _norm_runs(h), latin=latin, kr=kr, bold=True,
                      color=header_color, size=header_size,
                      align=(halign[j] if j < len(halign) else "left"))
    # 본문
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.cell(1 + i, j)
            if zebra and i % 2 == 1:
                shade_cell(c, zebra_fill)
            set_cell_margins(c, top=60, bottom=60, left=110, right=110)
            set_cell_runs(c, _norm_runs(val), latin=latin, kr=kr, size=font_size,
                          align=(aligns[j] if j < len(aligns) else "left"))
    if col_widths:
        for j, w in enumerate(col_widths):
            for r in range(1 + len(rows)):
                table.cell(r, j).width = Inches(w)
    return table


# ──────────────────────────────────────────────────────── 콜아웃 / 밴드 / KPI
def callout(doc, runs, accent=None, fill=None, bar_pt=3.0):
    """좌측 액센트 바 + surface 음영의 단일 셀 표(강조 박스)."""
    th = theme_of(doc)
    latin, kr = th.get("_latin"), th.get("_kr")
    accent = accent or th.get("primary") or "#333333"
    fill = fill or th.get("surface_2") or th.get("surface") or "#F2F2F2"
    table = doc.add_table(rows=1, cols=1)
    table.style = None
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    set_table_borders(table, fill, weight_pt=0.25, inside=False, outside=True)
    set_cell_border(cell, left={"val": "single", "sz": int(round(bar_pt * 8)), "color": accent})
    set_cell_runs(cell, _with_theme_fonts(doc, runs), latin=latin, kr=kr, vertical="center")
    return table


def band(doc, fill, accent=None):
    """전폭 색 밴드(표지/클로징용 단일 셀 표). 반환 cell 에 add_paragraph 로 내용 채움."""
    table = doc.add_table(rows=1, cols=1)
    table.style = None
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_margins(cell, top=320, bottom=320, left=260, right=260)
    set_table_borders(table, fill, weight_pt=0.25, inside=False, outside=True)
    if accent:
        set_cell_border(cell, left={"val": "single", "sz": 36, "color": accent})
    return cell


def kpi_strip(doc, items, value_color=None, label_color=None, value_size=22, label_size=9.5):
    """KPI 스트립: [(value, label), ...] 를 무테 표 한 행으로(값 크게·라벨 작게)."""
    th = theme_of(doc)
    latin, kr = th.get("_latin"), th.get("_kr")
    vcol = value_color or th.get("primary_text") or th.get("primary") or th.get("ink")
    lcol = label_color or th.get("muted")
    n = len(items)
    table = doc.add_table(rows=1, cols=n)
    table.style = None
    table.autofit = True
    set_table_borders(table, th.get("border") or "#CCCCCC", weight_pt=0.75,
                      inside=False, outside=False)
    for j, (value, label) in enumerate(items):
        c = table.cell(0, j)
        set_cell_margins(c, top=80, bottom=80, left=60, right=140)
        set_cell_runs(c, [(value, dict(latin=latin, kr=kr, size=value_size, bold=True, color=vcol))],
                      align="left", vertical="top")
        set_cell_runs(c, [(label, dict(latin=latin, kr=kr, size=label_size, color=lcol))],
                      align="left", first=False)
    return table


# ──────────────────────────────────────────────────────── 헤더/푸터 + 페이지번호
def _field(paragraph, instr, latin=None, color=None, size=None):
    """필드 코드(PAGE/NUMPAGES 등). Word 가 열 때 계산."""
    r = paragraph.add_run()
    set_run_font(r, latin=latin, color=color, size=size)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = f" {instr} "
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, sep, t, end):
        r._r.append(el)
    return r


def set_footer(doc, left=None, page_numbers=True, total_pages=True, color=None,
               size=8.5, hairline=True):
    """푸터: 좌측 텍스트 [tab] 우측 'N / M' 페이지번호 + 상단 헤어라인."""
    th = theme_of(doc)
    col = color or th.get("muted") or "#888888"
    latin, kr = th.get("_latin"), th.get("_kr")
    sec = doc.sections[0]
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    content_w = sec.page_width - sec.left_margin - sec.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(content_w, WD_TAB_ALIGNMENT.RIGHT)
    if hairline:
        _para_border(p, "top", th.get("border") or "#CCCCCC", weight_pt=0.75, space=4)
    if left:
        r = p.add_run(left)
        set_run_font(r, latin=latin, kr=kr, size=size, color=col)
    if page_numbers:
        r = p.add_run("\t")
        set_run_font(r, latin=latin, size=size, color=col)
        _field(p, "PAGE", latin, col, size)
        if total_pages:
            r = p.add_run(" / ")
            set_run_font(r, latin=latin, size=size, color=col)
            _field(p, "NUMPAGES", latin, col, size)
    return footer


def add_page_break(doc):
    doc.add_page_break()
